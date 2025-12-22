import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# -------------------------- 配置项 --------------------------
BASE_DIR = "results/translated_results"
ORIGINAL_IMG_DIR = "results/extracted_images"  # 原图目录
TRANS_IMG_DIR = os.path.join(BASE_DIR, "translated_images")
AUDIO_DIR = os.path.join(BASE_DIR, "audio_output")
OUTPUT_DOCX = "翻译汇总报告_链接版.docx"


# -----------------------------------------------------------

def add_hyperlink(paragraph, text, url):
    """为段落添加超链接的辅助函数"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                          is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set('{%s}id' % 'http://schemas.openxmlformats.org/officeDocument/2006/relationships', r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set('{%s}val' % 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set('{%s}val' % 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_word_report():
    doc = Document()
    doc.add_heading('图片翻译与语音合成报告', 0)

    # 获取原图并按名称排序
    img_files = [f for f in os.listdir(ORIGINAL_IMG_DIR) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
    img_files.sort()

    if not img_files:
        print(f"⚠️ 错误：在 {ORIGINAL_IMG_DIR} 文件夹中未找到图片。")
        return

    print(f"🚀 开始生成文档，共 {len(img_files)} 组数据...")

    for img_name in img_files:
        base_name = os.path.splitext(img_name)[0]

        # 路径匹配 (根据你的截图：image_xxx_translated.jpg)
        orig_path = os.path.join(ORIGINAL_IMG_DIR, img_name)
        trans_path = os.path.join(TRANS_IMG_DIR, f"{base_name}_translated.jpg")

        # 音频相对路径 (Word 相对于音频文件的路径)
        # 假设 Word 保存在 BASE_DIR 的同级目录，音频在 BASE_DIR/audio_output
        audio_rel_path = f"{BASE_DIR}/audio_output/{base_name}.mp3"

        # 1. 添加标题
        doc.add_heading(f"项目: {base_name}", level=1)

        # 2. 插入原图 (宽度设为 6 英寸，保证大图清晰)
        doc.add_paragraph("【原始图片】").bold = True
        doc.add_picture(orig_path, width=Inches(6.0))

        # 3. 插入翻译图
        doc.add_paragraph("\n【翻译结果】").bold = True
        if os.path.exists(trans_path):
            doc.add_picture(trans_path, width=Inches(6.0))
        else:
            doc.add_paragraph(f"(⚠️ 翻译图未找到: {base_name}_translated.jpg)")

        # 4. 添加音频超链接
        p_audio = doc.add_paragraph("\n🔊 ")
        if os.path.exists(os.path.join(AUDIO_DIR, f"{base_name}.mp3")):
            add_hyperlink(p_audio, "点击播放对应的合成语音 (MP3)", audio_rel_path)
        else:
            p_audio.add_run("(音频文件缺失)")

        # 5. 分页：每组内容占一页
        doc.add_page_break()

    doc.save(OUTPUT_DOCX)
    print(f"✨ 处理完成！文档已生成: {OUTPUT_DOCX}")


if __name__ == "__main__":
    create_base_dir = os.path.dirname(OUTPUT_DOCX)
    create_word_report()